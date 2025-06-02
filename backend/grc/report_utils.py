import os
import tempfile
import shutil
from django.db import connection
from .s3_functions import S3Client
from .models import Audit
import logging

# Set up logging
logger = logging.getLogger(__name__)

# Create Reports directory if it doesn't exist
REPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Reports')
os.makedirs(REPORTS_DIR, exist_ok=True)

def generate_and_upload_report(audit_id, user_id='system'):
    """
    Generate an audit report and upload it to S3, then delete the local file
    
    Args:
        audit_id: ID of the audit to generate a report for
        user_id: ID of the user requesting the report
        
    Returns:
        Dict containing result with success status and file info if successful
    """
    try:
        # Get audit details
        try:
            audit = Audit.objects.get(pk=audit_id)
            logger.info(f"Generating report for audit {audit_id} with status {audit.Status}")
        except Audit.DoesNotExist:
            logger.error(f"Audit with ID {audit_id} not found")
            return {'success': False, 'error': f"Audit with ID {audit_id} not found"}
        
        # Create a temporary directory for processing
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Generate a file name with audit details
            file_name = f"Audit_Report_{audit_id}_{audit.AuditType or 'Unknown'}.docx"
            temp_file_path = os.path.join(temp_dir, file_name)
            
            # Call the internal report generation function
            # This mimics what the API endpoint does but returns the file path rather than an HTTP response
            from .report_views import generate_report_file
            report_file = generate_report_file(audit_id, temp_file_path)
            
            if not report_file or not os.path.exists(temp_file_path):
                logger.error(f"Failed to generate report for audit {audit_id}")
                return {'success': False, 'error': "Failed to generate report"}
            
            # Copy the file to the Reports directory
            local_file_path = os.path.join(REPORTS_DIR, file_name)
            shutil.copy2(temp_file_path, local_file_path)
            logger.info(f"Report saved locally at {local_file_path}")
            
            # Initialize S3 client
            s3_client = S3Client()
            
            # Prepare metadata
            metadata = {
                'auditId': str(audit_id),
                'auditType': audit.AuditType or 'Unknown',
                'documentType': 'audit_report',
                'status': audit.Status or 'Unknown'
            }
            
            # Upload the file to S3
            upload_result = s3_client.upload_file(
                local_file_path, 
                user_id=user_id,
                **metadata  # Pass metadata as keyword arguments
            )
            
            if not isinstance(upload_result, dict) or 'error' in upload_result:
                error_msg = upload_result.get('error') if isinstance(upload_result, dict) else str(upload_result)
                logger.error(f"Failed to upload report to S3: {error_msg}")
                return {'success': False, 'error': f"Failed to upload report: {error_msg}"}
            
            logger.info(f"Report uploaded to S3 successfully: {upload_result}")
            
            # Delete the local file after successful upload
            try:
                os.remove(local_file_path)
                logger.info(f"Local report file deleted: {local_file_path}")
            except Exception as e:
                logger.warning(f"Failed to delete local report file: {str(e)}")
            
            return {
                'success': True,
                'message': 'Report generated and uploaded successfully',
                'file_info': upload_result.get('file'),
                'audit_id': audit_id
            }
            
        finally:
            # Clean up temporary directory
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    except Exception as e:
        logger.exception(f"Error generating and uploading report: {str(e)}")
        return {'success': False, 'error': str(e)}


def generate_report_file(audit_id, output_file_path, version=None):
    """
    Internal function to generate a report file for an audit
    
    Args:
        audit_id: ID of the audit to generate a report for
        output_file_path: Path to save the report to
        version: Optional specific version to generate
        
    Returns:
        Path to the generated file if successful, None otherwise
    """
    try:
        # Import here to avoid circular imports
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        import datetime
        
        # Get the audit details
        try:
            audit = Audit.objects.get(pk=audit_id)
        except Audit.DoesNotExist:
            logger.error(f"Audit with ID {audit_id} not found")
            return None
        
        # Get the findings (similar to report_views.py)
        findings = []
        if version and version.startswith('R'):
            # Get specific version data
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT 
                        av.ExtractedInfo,
                        av.Date
                    FROM 
                        audit_version av
                    WHERE 
                        av.AuditId = %s AND av.Version = %s
                """, [audit_id, version])
                
                version_row = cursor.fetchone()
                if not version_row:
                    logger.error(f"Version {version} not found for audit {audit_id}")
                    return None
                
                extracted_info = version_row[0]
                version_date = version_row[1]
                
                if extracted_info:
                    import json
                    # Parse JSON data
                    try:
                        version_data = json.loads(extracted_info)
                        
                        # Extract finding data from the version
                        for compliance_id, finding_data in version_data.items():
                            if compliance_id == '__metadata__' or compliance_id == 'overall_comments':
                                continue
                            
                            # Get compliance details
                            cursor.execute("""
                                SELECT 
                                    c.ComplianceItemDescription
                                FROM 
                                    compliance c
                                WHERE 
                                    c.ComplianceId = %s
                            """, [compliance_id])
                            
                            compliance_row = cursor.fetchone()
                            compliance_description = compliance_row[0] if compliance_row else "Unknown Compliance Item"
                            
                            # Convert version finding data
                            finding = {
                                'ComplianceId': compliance_id,
                                'MajorMinor': finding_data.get('major_minor', 'N/A'),
                                'Check': finding_data.get('check', 'N/A'),
                                'HowToVerify': finding_data.get('how_to_verify', 'N/A'),
                                'Evidence': finding_data.get('evidence', 'N/A'),
                                'DetailsOfFinding': finding_data.get('details_of_finding', 'N/A'),
                                'Impact': finding_data.get('impact', 'N/A'),
                                'Recommendation': finding_data.get('recommendation', 'N/A'),
                                'Comments': finding_data.get('comments', 'N/A'),
                                'CheckedDate': version_date,
                                'ComplianceItemDescription': compliance_description
                            }
                            findings.append(finding)
                            
                    except json.JSONDecodeError as e:
                        logger.error(f"Invalid JSON in ExtractedInfo: {str(e)}")
                        return None
        else:
            # Fetch findings directly from audit_findings table
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT 
                        af.AuditFindingsId,
                        af.ComplianceId,
                        af.MajorMinor,
                        af.Check,
                        af.HowToVerify,
                        af.Evidence,
                        af.DetailsOfFinding,
                        af.Impact,
                        af.Recommendation,
                        af.Comments,
                        af.CheckedDate,
                        c.ComplianceItemDescription
                    FROM 
                        audit_findings af
                    JOIN 
                        compliance c ON af.ComplianceId = c.ComplianceId
                    WHERE 
                        af.AuditId = %s
                    ORDER BY
                        af.AuditFindingsId
                """, [audit_id])
                
                columns = [col[0] for col in cursor.description]
                findings = [dict(zip(columns, row)) for row in cursor.fetchall()]
        
        # Check if we have findings
        if len(findings) == 0:
            logger.warning(f"No findings available for audit {audit_id}")
            return None
        
        # Create the document
        doc = Document()
        
        # Add title
        report_title = f'Audit Report - ID: {audit_id}'
        if version:
            report_title += f' (Version: {version})'
        title = doc.add_heading(report_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add report generation information
        report_date = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        doc.add_paragraph(f'Report generated on: {report_date}')
        
        # Add audit information section
        doc.add_heading('Audit Information', 1)
        audit_table = doc.add_table(rows=1, cols=2)
        audit_table.style = 'Table Grid'
        
        # Add header row
        header_cells = audit_table.rows[0].cells
        header_cells[0].text = 'Field'
        header_cells[1].text = 'Value'
        
        # Add audit details
        audit_data = [
            ('Audit ID', str(audit_id)),
            ('Status', getattr(audit, 'Status', 'Unknown')),
            ('Type', getattr(audit, 'AuditType', 'Unknown')),
            ('Due Date', getattr(audit, 'DueDate', 'N/A')),
            ('Review Status', getattr(audit, 'ReviewStatus', 'N/A')),
        ]
        
        if version:
            audit_data.append(('Report Version', version))
        
        for field, value in audit_data:
            row_cells = audit_table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = str(value)
        
        # Add a section for each finding
        doc.add_heading('Audit Findings', 1)
        
        for i, finding in enumerate(findings):
            # Create a heading for each finding
            doc.add_heading(f'Finding {i+1}: {finding.get("ComplianceItemDescription", "N/A")}', 2)
            
            # Create a table for the finding details
            finding_table = doc.add_table(rows=1, cols=2)
            finding_table.style = 'Table Grid'
            
            # Set column widths
            finding_table.autofit = False
            finding_table.columns[0].width = Inches(1.5)
            finding_table.columns[1].width = Inches(4.5)
            
            # Add headers
            header_cells = finding_table.rows[0].cells
            header_cells[0].text = 'ITEM'
            header_cells[1].text = 'DETAILS'
            
            # Add finding details
            finding_details = [
                ('Major/Minor', finding.get('MajorMinor', 'N/A')),
                ('Check', finding.get('Check', 'N/A')),
                ('How to Verify', finding.get('HowToVerify', 'N/A')),
                ('Evidence', finding.get('Evidence', 'N/A')),
                ('Details of Finding', finding.get('DetailsOfFinding', 'N/A')),
                ('Impact', finding.get('Impact', 'N/A')),
                ('Recommendation', finding.get('Recommendation', 'N/A')),
                ('Comments', finding.get('Comments', 'N/A')),
                ('Date Checked', finding.get('CheckedDate', 'N/A')),
            ]
            
            for field, value in finding_details:
                row_cells = finding_table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value) if value else 'N/A'
            
            # Add some space after each finding
            doc.add_paragraph()
        
        # Save the document
        doc.save(output_file_path)
        
        return output_file_path
        
    except Exception as e:
        logger.exception(f"Error generating report file: {str(e)}")
        return None 