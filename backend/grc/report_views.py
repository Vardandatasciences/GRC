from django.http import HttpResponse, FileResponse
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from django.db import connection
from .models import Audit
import os
import io
import tempfile
import shutil
from django.conf import settings
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

@api_view(['GET'])
def generate_audit_report(request, audit_id):
    """
    Generate and download an audit report in DOCX format with tables for each finding
    """
    try:
        print(f"DEBUG: generate_audit_report called for audit_id: {audit_id}")
        
        # Check if a specific version is requested
        version = request.query_params.get('version')
        print(f"DEBUG: Version parameter: {version}")
        
        # Create a temporary directory for the process
        temp_dir = tempfile.mkdtemp()
        output_file = os.path.join(temp_dir, f"audit_report_{audit_id}_{version if version else 'latest'}.docx")
        
        try:
            # Generate the report file
            report_file = generate_report_file(audit_id, output_file, version)
            
            if not report_file or not os.path.exists(output_file):
                print(f"ERROR: Failed to generate report for audit {audit_id}")
                return Response({"error": "Failed to generate report"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
            # Determine file name for download
            if version:
                download_filename = f"audit_report_{audit_id}_v{version}.docx"
            else:
                download_filename = f"audit_report_{audit_id}.docx"
            
            # Return the file for download
            response = FileResponse(
                open(output_file, 'rb'),
                as_attachment=True,
                filename=download_filename
            )
            return response
            
        finally:
            # Clean up temporary directory after sending response
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    except Exception as e:
        print(f"ERROR in generate_audit_report: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def generate_report_file(audit_id, output_file_path, version=None):
    """
    Internal function to generate a report file for an audit
    
    Args:
        audit_id: ID of the audit to generate a report for
        output_file_path: Path to save the report to
        version: Optional specific version to generate
        
    Returns:
        Path to the generated file if successful, None otherwise
    """
    try:
        # Get the audit details
        try:
            audit = Audit.objects.get(pk=audit_id)
            print(f"DEBUG: Found audit with ID {audit_id}, status: {getattr(audit, 'Status', 'Unknown')}")
        except Audit.DoesNotExist:
            print(f"ERROR: Audit with ID {audit_id} not found")
            return None
        
        # Get the findings
        findings = []
        if version and version.startswith('R'):
            # Get specific version data
            with connection.cursor() as cursor:
                print(f"DEBUG: Fetching version data for audit_id: {audit_id}, version: {version}")
                cursor.execute("""
                    SELECT 
                        av.ExtractedInfo,
                        av.Date
                    FROM 
                        audit_version av
                    WHERE 
                        av.AuditId = %s AND av.Version = %s
                """, [audit_id, version])
                
                version_row = cursor.fetchone()
                if not version_row:
                    print(f"ERROR: Version {version} not found for audit {audit_id}")
                    return None
                
                extracted_info = version_row[0]
                version_date = version_row[1]
                
                if extracted_info:
                    import json
                    # Parse JSON data
                    try:
                        version_data = json.loads(extracted_info)
                        print(f"DEBUG: Successfully parsed JSON data, found {len(version_data)} items")
                        
                        # Extract finding data from the version
                        for compliance_id, finding_data in version_data.items():
                            if compliance_id == '__metadata__' or compliance_id == 'overall_comments':
                                continue
                            
                            # Get compliance details
                            cursor.execute("""
                                SELECT 
                                    c.ComplianceItemDescription
                                FROM 
                                    compliance c
                                WHERE 
                                    c.ComplianceId = %s
                            """, [compliance_id])
                            
                            compliance_row = cursor.fetchone()
                            compliance_description = compliance_row[0] if compliance_row else "Unknown Compliance Item"
                            
                            # Convert version finding data
                            finding = {
                                'ComplianceId': compliance_id,
                                'MajorMinor': finding_data.get('major_minor', 'N/A'),
                                'Check': finding_data.get('check', 'N/A'),
                                'HowToVerify': finding_data.get('how_to_verify', 'N/A'),
                                'Evidence': finding_data.get('evidence', 'N/A'),
                                'DetailsOfFinding': finding_data.get('details_of_finding', 'N/A'),
                                'Impact': finding_data.get('impact', 'N/A'),
                                'Recommendation': finding_data.get('recommendation', 'N/A'),
                                'Comments': finding_data.get('comments', 'N/A'),
                                'CheckedDate': version_date,
                                'ComplianceItemDescription': compliance_description
                            }
                            findings.append(finding)
                            
                    except json.JSONDecodeError as e:
                        print(f"ERROR: Invalid JSON in ExtractedInfo: {str(e)}")
                        return None
        else:
            # Fetch findings directly from audit_findings table
            with connection.cursor() as cursor:
                print(f"DEBUG: Executing SQL query to fetch findings for audit_id: {audit_id}")
                cursor.execute("""
                    SELECT 
                        af.AuditFindingsId,
                        af.ComplianceId,
                        af.MajorMinor,
                        af.Check,
                        af.HowToVerify,
                        af.Evidence,
                        af.DetailsOfFinding,
                        af.Impact,
                        af.Recommendation,
                        af.Comments,
                        af.CheckedDate,
                        c.ComplianceItemDescription
                    FROM 
                        audit_findings af
                    JOIN 
                        compliance c ON af.ComplianceId = c.ComplianceId
                    WHERE 
                        af.AuditId = %s
                    ORDER BY
                        af.AuditFindingsId
                """, [audit_id])
                
                columns = [col[0] for col in cursor.description]
                findings = [dict(zip(columns, row)) for row in cursor.fetchall()]
                print(f"DEBUG: Query successful, found {len(findings)} findings")
        
        # Count the number of findings
        findings_count = len(findings)
        print(f"DEBUG: Found {findings_count} findings for audit {audit_id}")
        
        if findings_count == 0:
            print(f"DEBUG: No findings available for audit {audit_id}")
            return None
        
        # Create the document
        doc = Document()
        
        # Add title
        report_title = f'Audit Report - ID: {audit_id}'
        if version:
            report_title += f' (Version: {version})'
        title = doc.add_heading(report_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add report generation information
        report_date = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        doc.add_paragraph(f'Report generated on: {report_date}')
        
        # Add audit information section
        doc.add_heading('Audit Information', 1)
        audit_table = doc.add_table(rows=1, cols=2)
        audit_table.style = 'Table Grid'
        
        # Add header row
        header_cells = audit_table.rows[0].cells
        header_cells[0].text = 'Field'
        header_cells[1].text = 'Value'
        
        # Add audit details
        audit_data = [
            ('Audit ID', str(audit_id)),
            ('Status', getattr(audit, 'Status', 'Unknown')),
            ('Type', getattr(audit, 'AuditType', 'Unknown')),
            ('Due Date', getattr(audit, 'DueDate', 'N/A')),
            ('Review Status', getattr(audit, 'ReviewStatus', 'N/A')),
        ]
        
        if version:
            audit_data.append(('Report Version', version))
        
        for field, value in audit_data:
            row_cells = audit_table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = str(value)
        
        # Add a section for each finding
        doc.add_heading('Audit Findings', 1)
        
        for i, finding in enumerate(findings):
            # Create a heading for each finding
            doc.add_heading(f'Finding {i+1}: {finding.get("ComplianceItemDescription", "N/A")}', 2)
            
            # Create a table for the finding details
            finding_table = doc.add_table(rows=1, cols=2)
            finding_table.style = 'Table Grid'
            
            # Set column widths
            finding_table.autofit = False
            finding_table.columns[0].width = Inches(1.5)
            finding_table.columns[1].width = Inches(4.5)
            
            # Add headers
            header_cells = finding_table.rows[0].cells
            header_cells[0].text = 'ITEM'
            header_cells[1].text = 'DETAILS'
            
            # Add finding details
            finding_details = [
                ('Major/Minor', finding.get('MajorMinor', 'N/A')),
                ('Check', finding.get('Check', 'N/A')),
                ('How to Verify', finding.get('HowToVerify', 'N/A')),
                ('Evidence', finding.get('Evidence', 'N/A')),
                ('Details of Finding', finding.get('DetailsOfFinding', 'N/A')),
                ('Impact', finding.get('Impact', 'N/A')),
                ('Recommendation', finding.get('Recommendation', 'N/A')),
                ('Comments', finding.get('Comments', 'N/A')),
                ('Date Checked', finding.get('CheckedDate', 'N/A')),
            ]
            
            for field, value in finding_details:
                row_cells = finding_table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value) if value else 'N/A'
            
            # Add some space after each finding
            doc.add_paragraph()
        
        # Save the document
        doc.save(output_file_path)
        
        return output_file_path
        
    except Exception as e:
        print(f"ERROR generating report file: {str(e)}")
        import traceback
        traceback.print_exc()
        return None 